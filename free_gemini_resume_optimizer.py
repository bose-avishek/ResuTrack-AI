import os
import re
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Initialize the official Google GenAI client using your environment key
client = genai.Client()

def strip_markdown(text):
    """Safely cleans out markdown styling tags before rendering into Word."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    return text.strip()

def add_bottom_border(paragraph):
    """Injects a native XML bottom border straight into the Word paragraph element."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 6/8 pt line thickness
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')  # Pure black hex color
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_word_doc(md_path, word_path):
    """Converts the text file into an elegantly styled, Calibri Word document."""
    doc = Document()
    
    # Enforce strict 1-inch margins on all sides (Recruiter Standard)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Establish global base typography using Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)  # Explicitly default normal text to black

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
            
        # 1. Main Name / Document Title Header (# )
        if line.startswith("# "):
            p = doc.add_heading(strip_markdown(clean_line), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centered name look
            p.paragraph_format.space_after = Pt(2)  # Tight spacing leading into contact details
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Overrides default Word blue to black

        # 2. Main Section Categories (## ) like "Work Experience", "Skills"
        elif line.startswith("## "):
            p = doc.add_heading(strip_markdown(clean_line), level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True  # Prevents orphan category titles
            add_bottom_border(p)  # Adds the clean horizontal divider line
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Overrides default Word blue to black

        # 3. Subheadings (### ) like Company Name, Job Title, Location, Dates
        elif line.startswith("### "):
            p = doc.add_heading(strip_markdown(clean_line), level=3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True  # Automatically bolds the entire structural metadata line
                run.font.color.rgb = RGBColor(0, 0, 0)  # Overrides default Word blue to black

        # 4. Bullet Points (Checks for markdown symbols * / - OR literal unicode bullets like •)
        elif line.startswith("* ") or line.startswith("- ") or line.startswith("• "):
            bullet_content = clean_line[2:]
            p = doc.add_paragraph(strip_markdown(bullet_content), style='List Bullet')
            p.paragraph_format.space_after = Pt(4)     # Clean breathing room between points
            p.paragraph_format.line_spacing = 1.15     # Modern, clean line layout
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)

        # 5. Regular Paragraph blocks (e.g., Professional Summary or Contact Block)
        else:
            if clean_line.startswith("•") or clean_line.startswith("-") or clean_line.startswith("*"):
                p = doc.add_paragraph(strip_markdown(clean_line[1:].strip()), style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
            else:
                p = doc.add_paragraph(strip_markdown(clean_line))
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                
                # Center contact info lines
                if any(item in clean_line.lower() for item in ["@ ", "linkedin.com", "phone:", "+", " | "]):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(14)
                
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
    doc.save(word_path)

def pipeline_execution():
    os.makedirs("job_listings", exist_ok=True)
    os.makedirs("tailored_resumes", exist_ok=True)
    
    if not os.path.exists("master_resume.md"):
        print("Error: 'master_resume.md' missing. Please create it first.")
        return

    with open("master_resume.md", "r", encoding="utf-8") as f:
        master_content = f.read()

    system_prompt = (
        "You are an expert career coach and professional resume writer. "
        "Your task is to rewrite and optimize the user's master resume so that it is universally ATS-friendly "
        "and impactful for roles in the BFSI, Consulting, and Technology industries.\n\n"
        "Please follow these strict guidelines:\n"
        "1. Standardized Structure: Reorganize the content using universally recognized ATS headers "
        "with exactly two hash marks (e.g., '## Professional Summary', '## Skills', '## Work Experience', '## Education'). "
        "Ensure experience is in reverse-chronological order.\n"
        "2. Impact-Driven Bullets: Rewrite experience bullets using strong action verbs (e.g., spearheaded, optimized, drove). "
        "Focus heavily on quantifiable achievements and results rather than just listing daily responsibilities.\n"
        "3. Industry Keyword Integration: Identify essential keywords, methodologies, and hard skills standard for "
        "BFSI, Consulting, and Technology, and naturally integrate them throughout the summary and experience sections. "
        "Create a concise, comma-separated '## Skills' section.\n"
        "4. Clean Contact Info: Keep the email, phone, location, and professional links grouped on a single, pipe-separated line "
        "directly underneath the main candidate name.\n"
        "5. Experience Headers: Format job entries cleanly on their own structural line starting with three hash marks '### '. "
        "Include the Job Title, Company Name, Location, and Dates on this line separated by pipes so it stands out natively (e.g., '### Business Analyst | HSBC | Kolkata, India | Dates').\n"
        "6. Plain Text Output: Return only clean Markdown text. Use standard markdown symbols like '* ' or '- ' for list bullets. Do not use markdown bold symbols inside headings. "
        "Do not use tables, multiple columns, or complex graphic layouts. Do not include introductory or concluding conversational chat text.\n"
        "7. The XYZ Formula: Structure your experience bullet points to show impact following the Google formula: "
        "'Accomplished [X] as measured by [Y], by doing [Z].'"
    )

    for filename in os.listdir("job_listings"):
        if filename.endswith(".txt"):
            job_path = os.path.join("job_listings", filename)
            
            # Extract name and extension cleanly using tuple unpacking
            base_name_string, _ = os.path.splitext(filename)
            
            # Clean special characters out of the title string safely
            clean_file_title = re.sub(r'[\s\-]+', '_', base_name_string).strip('_')
            
            out_md = os.path.join("tailored_resumes", f"tailored_{clean_file_title}.md")
            out_docx = os.path.join("tailored_resumes", f"tailored_{clean_file_title}.docx")
            
            print(f"Optimizing: {filename} via Gemini API...")
            
            with open(job_path, "r", encoding="utf-8") as f:
                job_desc = f.read()
                
            try:
                response = client.models.generate_content(
                    model='gemini-3.6-flash',
                    contents=f"Master Resume:\n{master_content}\n\nJob Description:\n{job_desc}",
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                    ),
                )
                
                with open(out_md, "w", encoding="utf-8") as f:
                    f.write(response.text)
                    
                build_word_doc(out_md, out_docx)
                print(f"Success! Clean file generated: tailored_{clean_file_title}.docx")
                
            except Exception as e:
                print(f"Error handling file processing step: {e}")

if __name__ == "__main__":
    pipeline_execution()
